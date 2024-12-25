from selenium import webdriver
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.common.by import By
from selenium.webdriver.chrome.service import Service
from webdriver_manager.chrome import ChromeDriverManager
from docx import Document
import time

def search_and_extract(question_number, driver):
    # Navigate to Google
    driver.get("https://www.google.com")
    search_box = driver.find_element(By.NAME, 'q')

    # Enter the search query
    search_query = f"amazon aws certified solutions architect - associate saa-c03 question {question_number}"
    search_box.send_keys(search_query)
    search_box.send_keys(Keys.RETURN)
    
    # Wait for results to load and click the first link
    time.sleep(2)
    first_result = driver.find_element(By.CSS_SELECTOR, 'h3')
    first_result.click()

    # Wait for the page to load completely
    time.sleep(2)

    # Extract the question, choices, and correct answer
    question = driver.find_element(By.CLASS_NAME, 'card-text').text
    choices_elements = driver.find_elements(By.CLASS_NAME, 'multi-choice-item')
    choices = [choice.text for choice in choices_elements]
    correct_answer = driver.find_element(By.CLASS_NAME, 'correct-hidden').text

    return question, choices, correct_answer

def main():
    # Document setup
    doc = Document()
    doc.add_heading('AWS Certified Solutions Architect - Associate SAA-C03 Questions', level=1)
    
    # Selenium setup
    driver = webdriver.Chrome(service=Service(ChromeDriverManager().install()))

    try:
        for i in range(1, 11):  # Loop through the first 10 questions
            question, choices, correct_answer = search_and_extract(i, driver)
            doc.add_heading(f'Question {i}', level=2)
            doc.add_paragraph(question)
            doc.add_paragraph('Choices:')
            for idx, choice in enumerate(choices):
                doc.add_paragraph(f'{chr(65 + idx)}. {choice}', style='List Bullet')
            doc.add_paragraph(f'Correct Answer: {correct_answer}')
            doc.add_page_break()
    
    finally:
        driver.quit()
        doc.save('AWS_Certified_Solutions_Architect_Questions.docx')
        print("Document has been created and saved!")

if __name__ == '__main__':
    main()
